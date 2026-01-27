#!/usr/bin/env python3
# scripts/demo_legal_parser.py
"""
Demo script for VietnamLegalDocxParser.

Usage:
    python scripts/demo_legal_parser.py <path_to_docx>
    python scripts/demo_legal_parser.py --create-sample
    
This script demonstrates the parser's capabilities:
1. Parsing a Vietnamese legal document
2. Building hierarchical structure
3. Generating chunks with metadata
4. Displaying statistics
"""

import argparse
import json
import sys
from pathlib import Path

# Add parent directory to path for imports
sys.path.insert(0, str(Path(__file__).parent.parent))

from indexing.loaders.vietnam_legal_docx_parser import (
    VietnamLegalDocxParser,
    LegalNodeType,
    parse_vietnam_legal_docx,
)


def create_sample_docx(output_path: Path) -> None:
    """Create a sample Vietnamese legal document for testing."""
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("Error: python-docx is required. Install with: pip install python-docx")
        sys.exit(1)
    
    doc = Document()
    
    # Add boilerplate (should be filtered)
    p = doc.add_paragraph("QUỐC HỘI")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph("Độc lập - Tự do - Hạnh phúc")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph("-" * 50)
    
    # Law title
    p = doc.add_paragraph("LUẬT")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    p = doc.add_paragraph("GIAO DỊCH ĐIỆN TỬ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    
    doc.add_paragraph("Số: 20/2023/QH15")
    
    # Chapter I
    doc.add_paragraph("")
    p = doc.add_paragraph("Chương I. QUY ĐỊNH CHUNG")
    p.runs[0].bold = True
    
    # Article 1
    p = doc.add_paragraph("Điều 1. Phạm vi điều chỉnh")
    p.runs[0].bold = True
    doc.add_paragraph(
        "Luật này quy định về giao dịch điện tử; quyền, nghĩa vụ và trách nhiệm "
        "của cơ quan, tổ chức, cá nhân trong giao dịch điện tử; giá trị pháp lý "
        "của thông điệp dữ liệu, chữ ký điện tử, chứng thư điện tử; an toàn, "
        "an ninh trong giao dịch điện tử; giải quyết tranh chấp và xử lý vi phạm "
        "pháp luật trong giao dịch điện tử."
    )
    
    # Article 2
    p = doc.add_paragraph("Điều 2. Đối tượng áp dụng")
    p.runs[0].bold = True
    doc.add_paragraph(
        "1. Cơ quan, tổ chức, cá nhân trực tiếp tham gia hoặc có liên quan đến "
        "giao dịch điện tử."
    )
    doc.add_paragraph(
        "2. Cơ quan, tổ chức, cá nhân nước ngoài tham gia giao dịch điện tử "
        "tại Việt Nam."
    )
    
    # Article 3 - Definitions
    p = doc.add_paragraph("Điều 3. Giải thích từ ngữ")
    p.runs[0].bold = True
    doc.add_paragraph("Trong Luật này, các từ ngữ dưới đây được hiểu như sau:")
    doc.add_paragraph(
        "1. Giao dịch điện tử là giao dịch được thực hiện bằng phương tiện điện tử."
    )
    doc.add_paragraph(
        "2. Phương tiện điện tử là phương tiện hoạt động dựa trên công nghệ điện, "
        "điện tử, kỹ thuật số, từ tính, truyền dẫn không dây, quang học, điện từ "
        "hoặc công nghệ tương tự."
    )
    doc.add_paragraph(
        "3. Thông điệp dữ liệu là thông tin được tạo ra, gửi đi, nhận và lưu trữ "
        "bằng phương tiện điện tử."
    )
    doc.add_paragraph(
        "4. Chữ ký điện tử là chữ ký được tạo lập dưới dạng từ, chữ, số, ký hiệu, "
        "âm thanh hoặc các hình thức khác bằng phương tiện điện tử."
    )
    doc.add_paragraph(
        "5. Chữ ký số là một dạng chữ ký điện tử được tạo ra bằng sự biến đổi "
        "một thông điệp dữ liệu sử dụng hệ thống mật mã không đối xứng."
    )
    
    # Chapter II
    doc.add_paragraph("")
    p = doc.add_paragraph("Chương II. THÔNG ĐIỆP DỮ LIỆU")
    p.runs[0].bold = True
    
    # Section 1
    p = doc.add_paragraph("Mục 1. Quy định chung về thông điệp dữ liệu")
    p.runs[0].bold = True
    
    # Article 4
    p = doc.add_paragraph("Điều 4. Giá trị pháp lý của thông điệp dữ liệu")
    p.runs[0].bold = True
    doc.add_paragraph(
        "1. Thông điệp dữ liệu không bị phủ nhận giá trị pháp lý chỉ vì đó là "
        "thông điệp dữ liệu."
    )
    doc.add_paragraph("2. Thông điệp dữ liệu có giá trị như văn bản nếu:")
    doc.add_paragraph(
        "a) Thông tin chứa trong thông điệp dữ liệu có thể truy cập và sử dụng "
        "được để tham chiếu khi cần thiết;"
    )
    doc.add_paragraph(
        "b) Nội dung của thông điệp dữ liệu được bảo đảm toàn vẹn kể từ khi "
        "được khởi tạo lần đầu dưới dạng thông điệp dữ liệu hoàn chỉnh;"
    )
    doc.add_paragraph(
        "c) Thông tin trong thông điệp dữ liệu có thể được lưu trữ và truy cập "
        "khi cần thiết."
    )
    
    # Article 5
    p = doc.add_paragraph("Điều 5. Nguyên bản của thông điệp dữ liệu")
    p.runs[0].bold = True
    doc.add_paragraph("Thông điệp dữ liệu được coi là nguyên bản khi đáp ứng đủ "
                      "các điều kiện sau:")
    doc.add_paragraph(
        "1. Nội dung của thông điệp dữ liệu được bảo đảm toàn vẹn kể từ khi "
        "được khởi tạo lần đầu dưới dạng một thông điệp dữ liệu hoàn chỉnh."
    )
    doc.add_paragraph(
        "2. Nội dung của thông điệp dữ liệu có thể được trình bày khi cần thiết "
        "để thực hiện quyền và nghĩa vụ của các bên trong giao dịch điện tử."
    )
    
    # Section 2
    p = doc.add_paragraph("Mục 2. Gửi, nhận thông điệp dữ liệu")
    p.runs[0].bold = True
    
    # Article 6
    p = doc.add_paragraph("Điều 6. Gửi thông điệp dữ liệu")
    p.runs[0].bold = True
    doc.add_paragraph(
        "1. Việc gửi thông điệp dữ liệu hoàn thành khi thông điệp dữ liệu đó "
        "được gửi đến hệ thống thông tin nằm ngoài sự kiểm soát của người khởi tạo."
    )
    doc.add_paragraph(
        "2. Thời điểm gửi thông điệp dữ liệu là thời điểm thông điệp dữ liệu "
        "được gửi đến hệ thống thông tin được người nhận chỉ định."
    )
    
    # Save document
    doc.save(str(output_path))
    print(f"✓ Created sample document: {output_path}")


def print_tree(node, indent: int = 0) -> None:
    """Print the document tree structure."""
    prefix = "  " * indent
    type_name = node.node_type.value
    
    if node.title:
        print(f"{prefix}[{type_name}] {node.identifier}: {node.title[:50]}...")
    else:
        print(f"{prefix}[{type_name}] {node.identifier}")
    
    for child in node.children:
        print_tree(child, indent + 1)


def main():
    parser = argparse.ArgumentParser(
        description="Demo VietnamLegalDocxParser - Parse Vietnamese legal DOCX files"
    )
    parser.add_argument(
        "file",
        nargs="?",
        help="Path to DOCX file to parse"
    )
    parser.add_argument(
        "--create-sample",
        action="store_true",
        help="Create a sample DOCX file for testing"
    )
    parser.add_argument(
        "--output",
        "-o",
        default="sample_legal_doc.docx",
        help="Output path for sample file (default: sample_legal_doc.docx)"
    )
    parser.add_argument(
        "--show-tree",
        action="store_true",
        help="Show the document tree structure"
    )
    parser.add_argument(
        "--show-chunks",
        type=int,
        metavar="N",
        help="Show first N chunks"
    )
    parser.add_argument(
        "--export-json",
        help="Export chunks to JSON file"
    )
    parser.add_argument(
        "--token-threshold",
        type=int,
        default=800,
        help="Token threshold for splitting (default: 800)"
    )
    
    args = parser.parse_args()
    
    # Create sample if requested
    if args.create_sample:
        output_path = Path(args.output)
        create_sample_docx(output_path)
        
        if not args.file:
            args.file = str(output_path)
    
    if not args.file:
        parser.print_help()
        print("\nError: Please provide a DOCX file to parse or use --create-sample")
        sys.exit(1)
    
    file_path = Path(args.file)
    if not file_path.exists():
        print(f"Error: File not found: {file_path}")
        sys.exit(1)
    
    print(f"\n{'='*60}")
    print(f"Parsing: {file_path}")
    print(f"Token threshold: {args.token_threshold}")
    print(f"{'='*60}\n")
    
    # Parse the document
    result = parse_vietnam_legal_docx(
        file_path,
        token_threshold=args.token_threshold
    )
    
    if not result.success:
        print("✗ Parsing failed!")
        for error in result.errors:
            print(f"  Error: {error}")
        sys.exit(1)
    
    print("✓ Parsing successful!\n")
    
    # Print statistics
    print("Statistics:")
    print(f"  - Total lines: {result.statistics.get('total_lines', 0)}")
    print(f"  - Total chunks: {result.statistics.get('total_chunks', 0)}")
    print(f"  - Chapters: {result.statistics.get('chapters', 0)}")
    print(f"  - Articles: {result.statistics.get('articles', 0)}")
    print(f"  - Clauses: {result.statistics.get('clauses', 0)}")
    print(f"  - Points: {result.statistics.get('points', 0)}")
    
    if result.warnings:
        print("\nWarnings:")
        for warning in result.warnings:
            print(f"  - {warning}")
    
    # Show tree if requested
    if args.show_tree and result.tree:
        print(f"\n{'='*60}")
        print("Document Tree Structure:")
        print(f"{'='*60}\n")
        print_tree(result.tree)
    
    # Show chunks if requested
    if args.show_chunks and result.chunks:
        n = min(args.show_chunks, len(result.chunks))
        print(f"\n{'='*60}")
        print(f"First {n} Chunks:")
        print(f"{'='*60}\n")
        
        for i, chunk in enumerate(result.chunks[:n]):
            print(f"--- Chunk {i+1}/{n} ---")
            print(f"ID: {chunk.chunk_id}")
            print(f"Prefix: {chunk.embedding_prefix}")
            print(f"Content ({len(chunk.content)} chars):")
            content_preview = chunk.content[:300] + "..." if len(chunk.content) > 300 else chunk.content
            print(f"  {content_preview}")
            print(f"Metadata: {json.dumps(chunk.metadata, ensure_ascii=False, indent=2)}")
            print()
    
    # Export to JSON if requested
    if args.export_json:
        output_path = Path(args.export_json)
        chunks_data = [chunk.to_dict() for chunk in result.chunks]
        
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump({
                "statistics": result.statistics,
                "chunks": chunks_data,
            }, f, ensure_ascii=False, indent=2)
        
        print(f"\n✓ Exported {len(result.chunks)} chunks to: {output_path}")


if __name__ == "__main__":
    main()
