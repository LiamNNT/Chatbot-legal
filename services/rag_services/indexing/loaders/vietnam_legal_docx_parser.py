# indexing/loaders/vietnam_legal_docx_parser.py
"""
VietnamLegalDocxParser - Hierarchical parser for Vietnamese legal documents.

This parser handles .docx (and .doc via LibreOffice conversion) files containing
Vietnamese legal texts with structure: Luật > Chương > Mục > Điều > Khoản > Điểm.

Features:
- Automatic .doc to .docx conversion (requires LibreOffice)
- Boilerplate removal and text normalization
- Hierarchical chunking with token-based splitting
- Tree-based relationship tracking (parent, siblings, children)
- Special handling for "Giải thích từ ngữ" (definitions) articles
"""

from __future__ import annotations

import logging
import os
import re
import shutil
import subprocess
import tempfile
import unicodedata
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

logger = logging.getLogger(__name__)


# =============================================================================
# Constants and Regex Patterns
# =============================================================================

# Default token threshold for splitting
DEFAULT_TOKEN_THRESHOLD = 800

# Boilerplate patterns to remove (case-insensitive)
BOILERPLATE_PATTERNS = [
    r"^\s*QUỐC\s*HỘI\s*$",
    r"^\s*CỘNG\s*HÒA\s*XÃ\s*HỘI\s*CHỦ\s*NGHĨA\s*VIỆT\s*NAM\s*$",
    r"^\s*Độc\s*lập\s*[-–—]\s*Tự\s*do\s*[-–—]\s*Hạnh\s*phúc\s*$",
    r"^\s*-+\s*$",  # Separator lines
    r"^\s*_{3,}\s*$",  # Underscores
    r"^\s*\*{3,}\s*$",  # Asterisks
    r"^\s*Số:\s*\d+.*$",  # Document numbers
    r"^\s*Căn\s*cứ\s*Hiến\s*pháp.*$",
    r"^\s*Trang\s*\d+\s*$",  # Page numbers
    r"^\s*\d+\s*/\s*\d+\s*$",  # Page numbers like "1 / 10"
]

# Compile boilerplate patterns
BOILERPLATE_RE = [re.compile(p, re.IGNORECASE | re.UNICODE) for p in BOILERPLATE_PATTERNS]

# Vietnamese legal structure patterns (line-anchored)
# CHƯƠNG (Chapter) - Roman numerals
CHUONG_PATTERN = re.compile(
    r"^\s*Chương\s+([IVXLCDM]+)\s*[:\.]?\s*(.*)$",
    re.IGNORECASE | re.UNICODE | re.MULTILINE
)

# MỤC (Section) - Arabic numerals
MUC_PATTERN = re.compile(
    r"^\s*Mục\s+(\d+)\s*[:\.]?\s*(.*)$",
    re.IGNORECASE | re.UNICODE | re.MULTILINE
)

# ĐIỀU (Article) - Arabic numerals, optionally followed by letters
DIEU_PATTERN = re.compile(
    r"^\s*Điều\s+(\d+[A-Za-zÀ-ỹđĐ]*)\.\s*(.*)$",
    re.IGNORECASE | re.UNICODE | re.MULTILINE
)

# KHOẢN (Clause) - Numbered items like "1. ", "2. "
KHOAN_PATTERN = re.compile(
    r"^\s*(\d+)\.\s+(.*)$",
    re.UNICODE | re.MULTILINE
)

# Explicit KHOẢN marker (less common)
KHOAN_EXPLICIT_PATTERN = re.compile(
    r"^\s*Khoản\s+(\d+)\s*[:\.]?\s*(.*)$",
    re.IGNORECASE | re.UNICODE | re.MULTILINE
)

# ĐIỂM (Point) - Letters like "a) ", "b) ", "đ) "
DIEM_PATTERN = re.compile(
    r"^\s*([a-zđ])\)\s+(.*)$",
    re.UNICODE | re.MULTILINE
)

# Pattern to detect "Giải thích từ ngữ" (definitions article)
DEFINITION_ARTICLE_PATTERNS = [
    r"Trong\s+(?:Luật|Nghị\s+định|Thông\s+tư)\s+này[,:]?\s*(?:các\s+)?từ\s+ngữ.*(?:được\s+)?hiểu\s+như\s+sau",
    r"Giải\s+thích\s+từ\s+ngữ",
    r"Định\s+nghĩa",
]
DEFINITION_RE = [re.compile(p, re.IGNORECASE | re.UNICODE) for p in DEFINITION_ARTICLE_PATTERNS]

# Law ID extraction from filename or content
LAW_ID_PATTERN = re.compile(
    r"(\d+)[/\-](\d{4})[/\-]?(QH\d*|NĐ[-\s]?CP|TT[-\s]?[A-Z]*)?",
    re.IGNORECASE | re.UNICODE
)


# =============================================================================
# Enums and Data Classes
# =============================================================================

class LegalNodeType(Enum):
    """Types of nodes in the Vietnamese legal document hierarchy."""
    LAW = "LAW"
    CHAPTER = "CHUONG"
    SECTION = "MUC"
    ARTICLE = "DIEU"
    CLAUSE = "KHOAN"
    POINT = "DIEM"
    DEFINITION_ITEM = "DEFINITION_ITEM"


@dataclass
class LegalNode:
    """
    Represents a node in the legal document tree structure.
    """
    node_type: LegalNodeType
    identifier: str  # e.g., "I", "1", "Điều 1", "a"
    title: Optional[str] = None
    content: str = ""
    
    # Tree relationships
    parent: Optional[LegalNode] = field(default=None, repr=False)
    children: List[LegalNode] = field(default_factory=list, repr=False)
    prev_sibling: Optional[LegalNode] = field(default=None, repr=False)
    next_sibling: Optional[LegalNode] = field(default=None, repr=False)
    
    # Position tracking
    start_line: int = 0
    end_line: int = 0
    
    # Metadata
    is_definition_article: bool = False
    raw_lines: List[str] = field(default_factory=list)
    
    def get_full_id(self) -> str:
        """Generate a full hierarchical ID for this node."""
        parts = []
        node = self
        while node is not None:
            parts.append(f"{node.node_type.value}={node.identifier}")
            node = node.parent
        return ":".join(reversed(parts))
    
    def get_lineage(self) -> List[str]:
        """Get the lineage path from root to this node."""
        lineage = []
        node = self
        while node is not None:
            lineage.append(node.node_type.value)
            node = node.parent
        return list(reversed(lineage))
    
    def get_ancestors(self) -> Dict[str, str]:
        """Get all ancestors as a dict for metadata."""
        ancestors = {}
        node = self.parent
        while node is not None:
            if node.node_type == LegalNodeType.LAW:
                ancestors["law_id"] = node.identifier
                ancestors["law_name"] = node.title or ""
            elif node.node_type == LegalNodeType.CHAPTER:
                ancestors["chapter"] = f"Chương {node.identifier}"
                if node.title:
                    ancestors["chapter"] += f" {node.title}"
            elif node.node_type == LegalNodeType.SECTION:
                ancestors["section"] = f"Mục {node.identifier}"
                if node.title:
                    ancestors["section"] += f" {node.title}"
            elif node.node_type == LegalNodeType.ARTICLE:
                ancestors["article_id"] = f"Điều {node.identifier}"
                ancestors["article_title"] = node.title or ""
            elif node.node_type == LegalNodeType.CLAUSE:
                ancestors["clause_no"] = node.identifier
            node = node.parent
        return ancestors


@dataclass
class LegalChunk:
    """
    Represents a chunk of legal content ready for indexing.
    This is the output contract for the parser.
    """
    chunk_id: str
    content: str
    embedding_prefix: str
    metadata: Dict[str, Any]
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        return {
            "chunk_id": self.chunk_id,
            "content": self.content,
            "embedding_prefix": self.embedding_prefix,
            "metadata": self.metadata,
        }


@dataclass
class ParseResult:
    """Result of parsing a legal document."""
    success: bool
    chunks: List[LegalChunk] = field(default_factory=list)
    tree: Optional[LegalNode] = None
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    statistics: Dict[str, Any] = field(default_factory=dict)


# =============================================================================
# Token Estimation
# =============================================================================

def estimate_tokens(text: str) -> int:
    """
    Estimate the number of tokens in a text.
    Uses a simple heuristic: ~0.75 tokens per Vietnamese syllable/word.
    For more accuracy, can integrate tiktoken or similar.
    """
    if not text or not text.strip():
        return 0
    # Split by whitespace and punctuation
    words = re.split(r'\s+', text.strip())
    # Filter empty strings
    words = [w for w in words if w]
    if not words:
        return 0
    # Vietnamese typically has ~1.3 tokens per word for embedding models
    return int(len(words) * 1.3)


# =============================================================================
# Main Parser Class
# =============================================================================

class VietnamLegalDocxParser:
    """
    Parser for Vietnamese legal documents in DOCX format.
    
    Features:
    - Handles both .docx and .doc files (with LibreOffice conversion)
    - Removes boilerplate (headers, footers, page numbers)
    - Builds hierarchical tree structure
    - Creates chunks with proper metadata and relationships
    - Special handling for definition articles
    """
    
    def __init__(
        self,
        token_threshold: int = DEFAULT_TOKEN_THRESHOLD,
        libreoffice_path: Optional[str] = None,
    ):
        """
        Initialize the parser.
        
        Args:
            token_threshold: Maximum tokens per chunk before splitting
            libreoffice_path: Path to LibreOffice executable (auto-detected if None)
        """
        self.token_threshold = token_threshold
        self.libreoffice_path = libreoffice_path or self._find_libreoffice()
    
    def _find_libreoffice(self) -> Optional[str]:
        """Find LibreOffice executable on the system."""
        candidates = [
            "soffice",
            "libreoffice",
            "/usr/bin/soffice",
            "/usr/bin/libreoffice",
            "/Applications/LibreOffice.app/Contents/MacOS/soffice",
            r"C:\Program Files\LibreOffice\program\soffice.exe",
        ]
        for candidate in candidates:
            if shutil.which(candidate):
                return candidate
        return None
    
    def parse(
        self,
        file_path: Union[str, Path],
        law_id: Optional[str] = None,
        law_name: Optional[str] = None,
    ) -> ParseResult:
        """
        Parse a legal document and return chunks.
        
        Args:
            file_path: Path to the .docx or .doc file
            law_id: Override law ID (auto-extracted if None)
            law_name: Override law name (auto-extracted if None)
            
        Returns:
            ParseResult with chunks, tree structure, and statistics
        """
        file_path = Path(file_path)
        result = ParseResult(success=False)
        
        # Validate file exists
        if not file_path.exists():
            result.errors.append(f"File not found: {file_path}")
            return result
        
        # Handle file format
        suffix = file_path.suffix.lower()
        
        if suffix == ".doc":
            # Try to convert .doc to .docx
            converted_path = self._convert_doc_to_docx(file_path)
            if converted_path is None:
                result.errors.append(
                    "Unsupported format: .doc requires conversion to .docx "
                    "(install libreoffice or provide .docx)"
                )
                return result
            file_path = converted_path
            result.warnings.append(f"Converted .doc to .docx: {converted_path}")
        elif suffix != ".docx":
            result.errors.append(f"Unsupported file format: {suffix}")
            return result
        
        try:
            # Extract text from DOCX
            paragraphs = self._extract_paragraphs(file_path)
            
            # Normalize text
            normalized_lines = self._normalize_text(paragraphs)
            
            if not normalized_lines:
                result.errors.append("No content found after normalization")
                return result
            
            # Auto-detect law_id and law_name if not provided
            if law_id is None:
                law_id = self._extract_law_id(file_path.name, normalized_lines)
            if law_name is None:
                law_name = self._extract_law_name(normalized_lines)
            
            # Build document tree
            tree = self._build_tree(normalized_lines, law_id, law_name, file_path.name)
            result.tree = tree
            
            # Generate chunks from tree
            chunks = self._generate_chunks(tree, file_path.name)
            result.chunks = chunks
            
            # Calculate statistics
            result.statistics = {
                "total_lines": len(normalized_lines),
                "total_chunks": len(chunks),
                "chapters": sum(1 for c in tree.children if c.node_type == LegalNodeType.CHAPTER),
                "articles": self._count_nodes(tree, LegalNodeType.ARTICLE),
                "clauses": self._count_nodes(tree, LegalNodeType.CLAUSE),
                "points": self._count_nodes(tree, LegalNodeType.POINT),
            }
            
            result.success = True
            
        except Exception as e:
            logger.exception(f"Error parsing document: {e}")
            result.errors.append(f"Parse error: {str(e)}")
        
        return result
    
    def _convert_doc_to_docx(self, doc_path: Path) -> Optional[Path]:
        """
        Convert .doc to .docx using LibreOffice.
        
        Returns:
            Path to converted file, or None if conversion failed
        """
        if not self.libreoffice_path:
            logger.warning("LibreOffice not found, cannot convert .doc file")
            return None
        
        try:
            with tempfile.TemporaryDirectory() as tmp_dir:
                # Run LibreOffice in headless mode
                cmd = [
                    self.libreoffice_path,
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", tmp_dir,
                    str(doc_path),
                ]
                
                result = subprocess.run(
                    cmd,
                    capture_output=True,
                    text=True,
                    timeout=60,
                )
                
                if result.returncode != 0:
                    logger.error(f"LibreOffice conversion failed: {result.stderr}")
                    return None
                
                # Find the converted file
                converted_name = doc_path.stem + ".docx"
                converted_path = Path(tmp_dir) / converted_name
                
                if not converted_path.exists():
                    logger.error("Converted file not found")
                    return None
                
                # Move to a persistent location
                output_path = doc_path.parent / converted_name
                shutil.copy2(converted_path, output_path)
                
                return output_path
                
        except subprocess.TimeoutExpired:
            logger.error("LibreOffice conversion timed out")
            return None
        except Exception as e:
            logger.error(f"Error during conversion: {e}")
            return None
    
    def _extract_paragraphs(self, docx_path: Path) -> List[str]:
        """
        Extract paragraphs from a DOCX file.
        Excludes headers and footers.
        """
        try:
            from docx import Document
            from docx.opc.exceptions import PackageNotFoundError
        except ImportError:
            raise ImportError(
                "python-docx is required. Install with: pip install python-docx"
            )
        
        try:
            doc = Document(docx_path)
        except PackageNotFoundError:
            raise ValueError(f"Invalid or corrupted DOCX file: {docx_path}")
        
        paragraphs = []
        
        # Extract main body paragraphs (this excludes headers/footers by default)
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        # Also extract from tables if present
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            paragraphs.append(text)
        
        return paragraphs
    
    def _normalize_text(self, paragraphs: List[str]) -> List[str]:
        """
        Normalize text by removing boilerplate and cleaning whitespace.
        """
        normalized = []
        
        for para in paragraphs:
            # Normalize Unicode (NFC form for Vietnamese)
            text = unicodedata.normalize("NFC", para)
            
            # Normalize whitespace
            text = re.sub(r'\s+', ' ', text).strip()
            
            # Skip empty lines
            if not text:
                continue
            
            # Check against boilerplate patterns
            is_boilerplate = False
            for pattern in BOILERPLATE_RE:
                if pattern.match(text):
                    is_boilerplate = True
                    break
            
            if not is_boilerplate:
                normalized.append(text)
        
        return normalized
    
    def _extract_law_id(self, filename: str, lines: List[str]) -> str:
        """Extract law ID from filename or content."""
        # Try filename first
        match = LAW_ID_PATTERN.search(filename)
        if match:
            parts = [g for g in match.groups() if g]
            return "/".join(parts)
        
        # Try first few lines of content
        for line in lines[:20]:
            match = LAW_ID_PATTERN.search(line)
            if match:
                parts = [g for g in match.groups() if g]
                return "/".join(parts)
        
        # Fallback
        return "UNKNOWN"
    
    def _extract_law_name(self, lines: List[str]) -> str:
        """Extract law name from content."""
        # Look for "LUẬT" followed by the name
        for i, line in enumerate(lines[:30]):
            if re.match(r"^\s*LUẬT\s*$", line, re.IGNORECASE):
                # Next non-empty line might be the name
                for j in range(i + 1, min(i + 5, len(lines))):
                    if lines[j] and not re.match(r"^\s*Số:", lines[j]):
                        return lines[j]
            elif re.match(r"^\s*LUẬT\s+(.+)$", line, re.IGNORECASE):
                match = re.match(r"^\s*LUẬT\s+(.+)$", line, re.IGNORECASE)
                if match:
                    return f"Luật {match.group(1)}"
        
        return "Unknown Law"
    
    def _build_tree(
        self,
        lines: List[str],
        law_id: str,
        law_name: str,
        source_file: str,
    ) -> LegalNode:
        """
        Build a hierarchical tree structure from the document lines.
        """
        # Create root node (Law)
        root = LegalNode(
            node_type=LegalNodeType.LAW,
            identifier=law_id,
            title=law_name,
        )
        
        # Current context for building hierarchy
        current_chapter: Optional[LegalNode] = None
        current_section: Optional[LegalNode] = None
        current_article: Optional[LegalNode] = None
        current_clause: Optional[LegalNode] = None
        
        # Track siblings for linking
        last_chapter: Optional[LegalNode] = None
        last_section: Optional[LegalNode] = None
        last_article: Optional[LegalNode] = None
        last_clause: Optional[LegalNode] = None
        last_point: Optional[LegalNode] = None
        
        i = 0
        while i < len(lines):
            line = lines[i]
            
            # Check for CHƯƠNG (Chapter)
            match = CHUONG_PATTERN.match(line)
            if match:
                chapter_id, chapter_title = match.groups()
                chapter = LegalNode(
                    node_type=LegalNodeType.CHAPTER,
                    identifier=chapter_id.upper(),
                    title=chapter_title.strip() if chapter_title else None,
                    parent=root,
                    start_line=i,
                )
                root.children.append(chapter)
                
                # Link siblings
                if last_chapter:
                    last_chapter.next_sibling = chapter
                    chapter.prev_sibling = last_chapter
                last_chapter = chapter
                
                current_chapter = chapter
                current_section = None
                current_article = None
                last_section = None
                last_article = None
                last_clause = None
                last_point = None
                i += 1
                continue
            
            # Check for MỤC (Section)
            match = MUC_PATTERN.match(line)
            if match:
                section_id, section_title = match.groups()
                parent = current_chapter if current_chapter else root
                section = LegalNode(
                    node_type=LegalNodeType.SECTION,
                    identifier=section_id,
                    title=section_title.strip() if section_title else None,
                    parent=parent,
                    start_line=i,
                )
                parent.children.append(section)
                
                # Link siblings
                if last_section and last_section.parent == parent:
                    last_section.next_sibling = section
                    section.prev_sibling = last_section
                last_section = section
                
                current_section = section
                current_article = None
                last_article = None
                last_clause = None
                last_point = None
                i += 1
                continue
            
            # Check for ĐIỀU (Article)
            match = DIEU_PATTERN.match(line)
            if match:
                article_id, article_title = match.groups()
                parent = current_section if current_section else (current_chapter if current_chapter else root)
                
                article = LegalNode(
                    node_type=LegalNodeType.ARTICLE,
                    identifier=article_id,
                    title=article_title.strip() if article_title else None,
                    parent=parent,
                    start_line=i,
                )
                
                # Check if this is a definition article
                full_line = line + " " + (lines[i + 1] if i + 1 < len(lines) else "")
                for pattern in DEFINITION_RE:
                    if pattern.search(full_line):
                        article.is_definition_article = True
                        break
                
                parent.children.append(article)
                
                # Link siblings
                if last_article and last_article.parent == parent:
                    last_article.next_sibling = article
                    article.prev_sibling = last_article
                last_article = article
                
                current_article = article
                current_clause = None
                last_clause = None
                last_point = None
                i += 1
                continue
            
            # Check for KHOẢN (Clause) - only if we're in an article context
            if current_article:
                match = KHOAN_PATTERN.match(line) or KHOAN_EXPLICIT_PATTERN.match(line)
                if match:
                    clause_id, clause_content = match.groups()
                    clause = LegalNode(
                        node_type=LegalNodeType.CLAUSE,
                        identifier=clause_id,
                        content=clause_content.strip(),
                        parent=current_article,
                        start_line=i,
                    )
                    clause.raw_lines.append(line)
                    current_article.children.append(clause)
                    
                    # Link siblings
                    if last_clause and last_clause.parent == current_article:
                        last_clause.next_sibling = clause
                        clause.prev_sibling = last_clause
                    last_clause = clause
                    
                    current_clause = clause
                    last_point = None
                    i += 1
                    continue
            
            # Check for ĐIỂM (Point) - only if we're in a clause context
            if current_clause:
                match = DIEM_PATTERN.match(line)
                if match:
                    point_id, point_content = match.groups()
                    point = LegalNode(
                        node_type=LegalNodeType.POINT,
                        identifier=point_id,
                        content=point_content.strip(),
                        parent=current_clause,
                        start_line=i,
                    )
                    point.raw_lines.append(line)
                    current_clause.children.append(point)
                    
                    # Link siblings
                    if last_point and last_point.parent == current_clause:
                        last_point.next_sibling = point
                        point.prev_sibling = last_point
                    last_point = point
                    
                    i += 1
                    continue
            
            # If no pattern matched, add to current context's content
            if current_clause:
                current_clause.content += " " + line
                current_clause.raw_lines.append(line)
            elif current_article:
                current_article.content += " " + line
                current_article.raw_lines.append(line)
            elif current_section:
                current_section.content += " " + line
            elif current_chapter:
                current_chapter.content += " " + line
            else:
                root.content += " " + line
            
            i += 1
        
        # Clean up content
        self._clean_node_content(root)
        
        return root
    
    def _clean_node_content(self, node: LegalNode) -> None:
        """Recursively clean content of all nodes."""
        node.content = node.content.strip()
        for child in node.children:
            self._clean_node_content(child)
    
    def _generate_chunks(self, tree: LegalNode, source_file: str) -> List[LegalChunk]:
        """
        Generate chunks from the document tree.
        Applies token-based splitting as needed.
        """
        chunks: List[LegalChunk] = []
        
        # Process all articles
        for article in self._find_all_nodes(tree, LegalNodeType.ARTICLE):
            article_chunks = self._chunk_article(article, source_file)
            chunks.extend(article_chunks)
        
        return chunks
    
    def _chunk_article(self, article: LegalNode, source_file: str) -> List[LegalChunk]:
        """
        Generate chunks for a single article.
        Applies splitting strategy based on token count and article type.
        """
        chunks: List[LegalChunk] = []
        
        # Get full article content
        article_content = self._get_node_full_content(article)
        article_tokens = estimate_tokens(article_content)
        
        # Special handling for definition articles
        if article.is_definition_article:
            definition_chunks = self._chunk_definition_article(article, source_file)
            if definition_chunks:
                return definition_chunks
        
        # If article fits within threshold, create single chunk
        if article_tokens <= self.token_threshold:
            chunk = self._create_chunk(article, article_content, source_file)
            return [chunk]
        
        # Article is too large, split by clauses
        if article.children:
            for clause in article.children:
                if clause.node_type == LegalNodeType.CLAUSE:
                    clause_chunks = self._chunk_clause(clause, article, source_file)
                    chunks.extend(clause_chunks)
        else:
            # No clauses, create single chunk anyway (with warning)
            logger.warning(
                f"Article {article.identifier} exceeds token threshold "
                f"({article_tokens} > {self.token_threshold}) but has no clauses to split"
            )
            chunk = self._create_chunk(article, article_content, source_file)
            chunks.append(chunk)
        
        return chunks
    
    def _chunk_clause(
        self,
        clause: LegalNode,
        article: LegalNode,
        source_file: str,
    ) -> List[LegalChunk]:
        """Generate chunks for a clause, splitting by points if needed."""
        chunks: List[LegalChunk] = []
        
        clause_content = self._get_node_full_content(clause)
        clause_tokens = estimate_tokens(clause_content)
        
        # If clause fits within threshold, create single chunk
        if clause_tokens <= self.token_threshold:
            chunk = self._create_chunk(clause, clause_content, source_file)
            return [chunk]
        
        # Clause is too large, split by points
        if clause.children:
            for point in clause.children:
                if point.node_type == LegalNodeType.POINT:
                    point_content = self._get_node_full_content(point)
                    chunk = self._create_chunk(point, point_content, source_file)
                    chunks.append(chunk)
        else:
            # No points, create single chunk anyway
            logger.warning(
                f"Clause {clause.identifier} in Article {article.identifier} "
                f"exceeds token threshold ({clause_tokens} > {self.token_threshold}) "
                f"but has no points to split"
            )
            chunk = self._create_chunk(clause, clause_content, source_file)
            chunks.append(chunk)
        
        return chunks
    
    def _chunk_definition_article(
        self,
        article: LegalNode,
        source_file: str,
    ) -> List[LegalChunk]:
        """
        Special handling for definition articles.
        Split into one chunk per definition item.
        """
        chunks: List[LegalChunk] = []
        
        # Definitions are typically in clauses numbered 1., 2., etc.
        for clause in article.children:
            if clause.node_type == LegalNodeType.CLAUSE:
                # Create a definition item node
                def_item = LegalNode(
                    node_type=LegalNodeType.DEFINITION_ITEM,
                    identifier=clause.identifier,
                    content=clause.content,
                    parent=article,
                    prev_sibling=clause.prev_sibling,
                    next_sibling=clause.next_sibling,
                )
                
                content = self._get_node_full_content(clause)
                chunk = self._create_chunk(def_item, content, source_file)
                chunks.append(chunk)
        
        return chunks
    
    def _create_chunk(
        self,
        node: LegalNode,
        content: str,
        source_file: str,
    ) -> LegalChunk:
        """Create a LegalChunk from a node."""
        # Build chunk_id
        chunk_id = node.get_full_id()
        
        # Build embedding prefix
        embedding_parts = []
        ancestors = node.get_ancestors()
        if ancestors.get("law_id"):
            embedding_parts.append(f"LAW={ancestors['law_id']}")
        if ancestors.get("chapter"):
            # Extract just the chapter number
            chapter_match = re.search(r"Chương\s+([IVXLCDM]+)", ancestors["chapter"])
            if chapter_match:
                embedding_parts.append(f"CHUONG={chapter_match.group(1)}")
        if ancestors.get("section"):
            section_match = re.search(r"Mục\s+(\d+)", ancestors["section"])
            if section_match:
                embedding_parts.append(f"MUC={section_match.group(1)}")
        if ancestors.get("article_id"):
            article_match = re.search(r"Điều\s+(\d+[A-Za-zÀ-ỹđĐ]*)", ancestors["article_id"])
            if article_match:
                embedding_parts.append(f"DIEU={article_match.group(1)}")
        if ancestors.get("clause_no"):
            embedding_parts.append(f"KHOAN={ancestors['clause_no']}")
        if node.node_type == LegalNodeType.POINT:
            embedding_parts.append(f"DIEM={node.identifier}")
        elif node.node_type == LegalNodeType.CLAUSE:
            embedding_parts.append(f"KHOAN={node.identifier}")
        elif node.node_type == LegalNodeType.DEFINITION_ITEM:
            embedding_parts.append(f"KHOAN={node.identifier}")
        
        embedding_prefix = " | ".join(embedding_parts)
        
        # Build metadata
        metadata = {
            "law_id": ancestors.get("law_id", ""),
            "law_name": ancestors.get("law_name", ""),
            "chapter": ancestors.get("chapter"),
            "section": ancestors.get("section"),
            "article_id": ancestors.get("article_id"),
            "article_title": ancestors.get("article_title"),
            "clause_no": ancestors.get("clause_no") or (
                node.identifier if node.node_type in (LegalNodeType.CLAUSE, LegalNodeType.DEFINITION_ITEM) else None
            ),
            "point_no": node.identifier if node.node_type == LegalNodeType.POINT else None,
            "parent_id": node.parent.get_full_id() if node.parent else None,
            "prev_sibling_id": node.prev_sibling.get_full_id() if node.prev_sibling else None,
            "next_sibling_id": node.next_sibling.get_full_id() if node.next_sibling else None,
            "source_file": source_file,
            "lineage": node.get_lineage(),
            "is_definition": node.node_type == LegalNodeType.DEFINITION_ITEM,
        }
        
        # Remove None values for cleaner output
        metadata = {k: v for k, v in metadata.items() if v is not None}
        
        return LegalChunk(
            chunk_id=chunk_id,
            content=content.strip(),
            embedding_prefix=embedding_prefix,
            metadata=metadata,
        )
    
    def _get_node_full_content(self, node: LegalNode) -> str:
        """Get the full content of a node including its children."""
        parts = []
        
        # Add node's own content
        if node.node_type == LegalNodeType.ARTICLE:
            parts.append(f"Điều {node.identifier}. {node.title or ''}")
        elif node.node_type == LegalNodeType.CLAUSE:
            parts.append(f"{node.identifier}. {node.content}")
        elif node.node_type == LegalNodeType.POINT:
            parts.append(f"{node.identifier}) {node.content}")
        else:
            if node.content:
                parts.append(node.content)
        
        # Add children content for articles
        if node.node_type == LegalNodeType.ARTICLE:
            if node.content:
                parts.append(node.content)
            for child in node.children:
                parts.append(self._get_node_full_content(child))
        elif node.node_type == LegalNodeType.CLAUSE:
            for child in node.children:
                parts.append(self._get_node_full_content(child))
        
        return " ".join(parts).strip()
    
    def _find_all_nodes(self, root: LegalNode, node_type: LegalNodeType) -> List[LegalNode]:
        """Find all nodes of a given type in the tree."""
        results = []
        
        def _traverse(node: LegalNode):
            if node.node_type == node_type:
                results.append(node)
            for child in node.children:
                _traverse(child)
        
        _traverse(root)
        return results
    
    def _count_nodes(self, root: LegalNode, node_type: LegalNodeType) -> int:
        """Count all nodes of a given type in the tree."""
        return len(self._find_all_nodes(root, node_type))


# =============================================================================
# Utility Functions
# =============================================================================

def parse_vietnam_legal_docx(
    file_path: Union[str, Path],
    law_id: Optional[str] = None,
    law_name: Optional[str] = None,
    token_threshold: int = DEFAULT_TOKEN_THRESHOLD,
) -> ParseResult:
    """
    Convenience function to parse a Vietnamese legal document.
    
    Args:
        file_path: Path to the .docx or .doc file
        law_id: Override law ID (auto-extracted if None)
        law_name: Override law name (auto-extracted if None)
        token_threshold: Maximum tokens per chunk before splitting
        
    Returns:
        ParseResult with chunks, tree structure, and statistics
    """
    parser = VietnamLegalDocxParser(token_threshold=token_threshold)
    return parser.parse(file_path, law_id=law_id, law_name=law_name)


if __name__ == "__main__":
    # Simple test
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python vietnam_legal_docx_parser.py <file.docx>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    result = parse_vietnam_legal_docx(file_path)
    
    if result.success:
        print(f"✓ Parsed successfully!")
        print(f"  Total chunks: {result.statistics.get('total_chunks', 0)}")
        print(f"  Chapters: {result.statistics.get('chapters', 0)}")
        print(f"  Articles: {result.statistics.get('articles', 0)}")
        print(f"  Clauses: {result.statistics.get('clauses', 0)}")
        print(f"  Points: {result.statistics.get('points', 0)}")
        
        print("\nFirst 3 chunks:")
        for chunk in result.chunks[:3]:
            print(f"\n--- {chunk.chunk_id} ---")
            print(f"Prefix: {chunk.embedding_prefix}")
            print(f"Content: {chunk.content[:200]}...")
    else:
        print(f"✗ Parse failed!")
        for error in result.errors:
            print(f"  Error: {error}")
