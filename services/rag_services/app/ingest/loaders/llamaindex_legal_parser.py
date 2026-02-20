# indexing/loaders/llamaindex_legal_parser.py
"""
LlamaIndex-based parser for Vietnamese legal documents.

This parser replaces the legacy regex-based VietnamLegalDocxParser with a more
robust LlamaIndex-powered implementation.

Features:
- LlamaParse for document parsing (PDF, DOCX)
- MarkdownElementNodeParser for hierarchical structure extraction
- SemanticSplitterNodeParser for intelligent chunking
- tiktoken for accurate token counting
- Support for tables and complex layouts

Usage:
    from app.ingest.loaders.llamaindex_legal_parser import LlamaIndexLegalParser
    
    parser = LlamaIndexLegalParser.from_env()
    result = await parser.parse("document.docx")
"""

from __future__ import annotations

import asyncio
import hashlib
import logging
import os
import re
import unicodedata
from dataclasses import dataclass, field
from enum import Enum
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple, Union

from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


# =============================================================================
# Configuration
# =============================================================================

class ParserConfig(BaseModel):
    """Configuration for LlamaIndex legal parser."""
    
    # LlamaParse settings
    llama_cloud_api_key: Optional[str] = Field(default=None)
    use_gpt4o_mode: bool = Field(default=True, description="Use GPT-4o for complex docs")
    
    # Parsing instruction for Vietnamese legal documents
    parsing_instruction: str = Field(default="""
Đây là văn bản pháp luật Việt Nam (Luật, Nghị định, Thông tư).
Hãy trích xuất và bảo tồn đầy đủ cấu trúc:

1. CẤU TRÚC PHÂN CẤP:
   - Chương (Chapter): "Chương I", "Chương II"...
   - Mục (Section): "Mục 1", "Mục 2"...  
   - Điều (Article): "Điều 1", "Điều 2"...
   - Khoản (Clause): "1.", "2.", "3."...
   - Điểm (Point): "a)", "b)", "c)"...

2. BẢNG BIỂU: Giữ nguyên format Markdown
3. GHI CHÚ: Các điều khoản sửa đổi, bổ sung
4. LOẠI BỎ: Header, footer, số trang, watermark

Output format: Markdown với cấu trúc rõ ràng.
""")
    
    # Chunking settings
    chunk_size: int = Field(default=800, description="Target tokens per chunk")
    chunk_overlap: int = Field(default=100, description="Overlap tokens between chunks")
    
    # Embedding settings
    embedding_model: str = Field(
        default="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
        description="Model for semantic chunking"
    )
    
    # LLM settings for GPT-4o mode
    llm_api_key: Optional[str] = Field(default=None)
    llm_base_url: Optional[str] = Field(default=None)
    
    # Fallback settings
    use_fallback_parser: bool = Field(default=True)
    
    @classmethod
    def from_env(cls) -> "ParserConfig":
        """Load configuration from environment variables."""
        from dotenv import load_dotenv
        load_dotenv()
        
        return cls(
            llama_cloud_api_key=os.getenv("LLAMA_CLOUD_API_KEY"),
            use_gpt4o_mode=os.getenv("LLAMA_PARSE_GPT4O_MODE", "true").lower() == "true",
            llm_api_key=os.getenv("OPENROUTER_API_KEY") or os.getenv("OPENAI_API_KEY"),
            llm_base_url=os.getenv("OPENAI_BASE_URL"),
            chunk_size=int(os.getenv("CHUNK_SIZE", "800")),
            chunk_overlap=int(os.getenv("CHUNK_OVERLAP", "100")),
        )


# =============================================================================
# Domain Models
# =============================================================================

class LegalNodeType(str, Enum):
    """Types of nodes in Vietnamese legal document hierarchy."""
    LAW = "LAW"
    CHAPTER = "CHUONG"
    SECTION = "MUC"
    ARTICLE = "DIEU"
    CLAUSE = "KHOAN"
    POINT = "DIEM"
    TABLE = "TABLE"
    DEFINITION = "DEFINITION"


@dataclass
class LegalChunk:
    """
    Represents a chunk of legal content ready for indexing.
    Compatible with the legacy VietnamLegalDocxParser output.
    """
    chunk_id: str
    content: str
    embedding_prefix: str
    metadata: Dict[str, Any]
    
    # Additional fields for LlamaIndex integration
    node_id: Optional[str] = None
    tokens: int = 0
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        return {
            "chunk_id": self.chunk_id,
            "content": self.content,
            "embedding_prefix": self.embedding_prefix,
            "metadata": self.metadata,
            "tokens": self.tokens,
        }


@dataclass
class ParseResult:
    """Result of parsing a legal document."""
    success: bool
    chunks: List[LegalChunk] = field(default_factory=list)
    errors: List[str] = field(default_factory=list)
    warnings: List[str] = field(default_factory=list)
    statistics: Dict[str, Any] = field(default_factory=dict)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    # Raw parsed content
    raw_content: str = ""
    tables: List[Dict[str, Any]] = field(default_factory=list)


# =============================================================================
# Token Utilities
# =============================================================================

def get_tokenizer():
    """Get tiktoken tokenizer for accurate token counting."""
    try:
        import tiktoken
        return tiktoken.get_encoding("cl100k_base")  # GPT-4 encoding
    except ImportError:
        logger.warning("tiktoken not installed, using fallback token estimation")
        return None


def count_tokens(text: str, tokenizer=None) -> int:
    """Count tokens accurately using tiktoken or fallback heuristic."""
    if not text:
        return 0
    
    if tokenizer is None:
        tokenizer = get_tokenizer()
    
    if tokenizer is not None:
        return len(tokenizer.encode(text))
    
    # Fallback: Vietnamese has roughly 1.5 tokens per word for multilingual models
    words = re.split(r'\s+', text.strip())
    return int(len([w for w in words if w]) * 1.5)


# =============================================================================
# Vietnamese Legal Structure Patterns
# =============================================================================

# More robust patterns than legacy parser
PATTERNS = {
    # Chapter: "Chương I", "Chương II", "CHƯƠNG I"
    "chapter": re.compile(
        r"^(?:Chương|CHƯƠNG)\s+([IVXLCDM]+)\.?\s*(.*)$",
        re.IGNORECASE | re.MULTILINE
    ),
    # Section: "Mục 1", "Mục 2"
    "section": re.compile(
        r"^(?:Mục|MỤC)\s+(\d+)\.?\s*(.*)$",
        re.IGNORECASE | re.MULTILINE
    ),
    # Article: "Điều 1.", "Điều 10a."
    "article": re.compile(
        r"^(?:Điều|ĐIỀU)\s+(\d+[a-zA-Z]?)\.?\s*(.*)$",
        re.IGNORECASE | re.MULTILINE
    ),
    # Clause: "1. ", "2. " - with validation for context
    # Only match if the number is small (1-99) to avoid false positives
    "clause": re.compile(
        r"^(\d{1,2})\.\s+(.*)$",
        re.MULTILINE
    ),
    # Point: "a)", "b)", "đ)" - including Vietnamese letters
    "point": re.compile(
        r"^([a-zđA-ZĐăâêôơưĂÂÊÔƠƯ])\)\s+(.*)$",
        re.IGNORECASE | re.MULTILINE
    ),
}

# Definition article patterns
DEFINITION_PATTERNS = [
    re.compile(r"Giải\s+thích\s+từ\s+ngữ", re.IGNORECASE),
    re.compile(r"Trong\s+(?:Luật|Nghị\s+định|Thông\s+tư)\s+này.*hiểu\s+như\s+sau", re.IGNORECASE),
    re.compile(r"Định\s+nghĩa", re.IGNORECASE),
]

# Document type patterns
DOC_TYPE_PATTERNS = {
    "LAW": [
        re.compile(r"^\s*LUẬT\s*$", re.IGNORECASE | re.MULTILINE),
        re.compile(r"QH\d*", re.IGNORECASE),
    ],
    "DECREE": [
        re.compile(r"^\s*NGHỊ\s*ĐỊNH\s*$", re.IGNORECASE | re.MULTILINE),
        re.compile(r"NĐ[-\s]?CP", re.IGNORECASE),
    ],
    "CIRCULAR": [
        re.compile(r"^\s*THÔNG\s*TƯ\s*$", re.IGNORECASE | re.MULTILINE),
        re.compile(r"TT[-\s]?[A-Z]+", re.IGNORECASE),
    ],
}

# Law ID pattern
LAW_ID_PATTERN = re.compile(
    r"(\d+)[/\-](\d{4})[/\-]?(QH\d*|NĐ[-\s]?CP|TT[-\s]?[A-Z]*)?",
    re.IGNORECASE
)


# =============================================================================
# Main Parser Class
# =============================================================================

class LlamaIndexLegalParser:
    """
    LlamaIndex-based parser for Vietnamese legal documents.
    
    Replaces the legacy VietnamLegalDocxParser with:
    - LlamaParse for document extraction
    - MarkdownElementNodeParser for structure preservation
    - SemanticSplitterNodeParser for intelligent chunking
    - tiktoken for accurate token counting
    """
    
    def __init__(self, config: ParserConfig):
        """Initialize the parser with configuration."""
        self.config = config
        self._llama_parser = None
        self._tokenizer = get_tokenizer()
        self._embed_model = None
    
    @classmethod
    def from_env(cls) -> "LlamaIndexLegalParser":
        """Create parser from environment variables."""
        return cls(ParserConfig.from_env())
    
    # =========================================================================
    # LlamaParse Integration
    # =========================================================================
    
    def _get_llama_parser(self):
        """Lazy initialization of LlamaParse."""
        if self._llama_parser is None:
            if not self.config.llama_cloud_api_key:
                raise ValueError(
                    "LLAMA_CLOUD_API_KEY is required. "
                    "Get your key from https://cloud.llamaindex.ai/"
                )
            
            try:
                from llama_parse import LlamaParse
                
                parser_kwargs = {
                    "api_key": self.config.llama_cloud_api_key,
                    "result_type": "markdown",
                    "parsing_instruction": self.config.parsing_instruction,
                    "verbose": False,
                    "language": "vi",
                    "skip_diagonal_text": True,  # Skip watermarks
                    "do_not_unroll_columns": False,
                }
                
                # Enable GPT-4o mode for complex documents
                if self.config.use_gpt4o_mode and self.config.llm_api_key:
                    parser_kwargs["gpt4o_mode"] = True
                    parser_kwargs["gpt4o_api_key"] = self.config.llm_api_key
                    logger.info("LlamaParse GPT-4o mode enabled")
                
                self._llama_parser = LlamaParse(**parser_kwargs)
                logger.info("LlamaParse initialized successfully")
                
            except ImportError:
                raise ImportError(
                    "llama-parse not installed. Run: pip install llama-parse>=0.5.0"
                )
        
        return self._llama_parser
    
    def _get_embed_model(self):
        """Get embedding model for semantic chunking."""
        if self._embed_model is None:
            try:
                from llama_index.embeddings.huggingface import HuggingFaceEmbedding
                self._embed_model = HuggingFaceEmbedding(
                    model_name=self.config.embedding_model
                )
            except ImportError:
                logger.warning("HuggingFaceEmbedding not available, semantic chunking disabled")
                self._embed_model = None
        return self._embed_model
    
    # =========================================================================
    # Main Parse Methods
    # =========================================================================
    
    async def parse(
        self,
        file_path: Union[str, Path],
        law_id: Optional[str] = None,
        law_name: Optional[str] = None,
    ) -> ParseResult:
        """
        Parse a legal document (PDF or DOCX).
        
        Args:
            file_path: Path to the document
            law_id: Override law ID (auto-detected if None)
            law_name: Override law name (auto-detected if None)
            
        Returns:
            ParseResult with chunks, metadata, and statistics
        """
        file_path = Path(file_path)
        result = ParseResult(success=False)
        
        # Validate file
        if not file_path.exists():
            result.errors.append(f"File not found: {file_path}")
            return result
        
        suffix = file_path.suffix.lower()
        
        try:
            if suffix == ".pdf":
                result = await self._parse_pdf(file_path, law_id, law_name)
            elif suffix in (".docx", ".doc"):
                result = await self._parse_docx(file_path, law_id, law_name)
            else:
                result.errors.append(f"Unsupported file format: {suffix}")
                return result
                
        except Exception as e:
            logger.exception(f"Error parsing document: {e}")
            result.errors.append(f"Parse error: {str(e)}")
        
        return result
    
    def parse_sync(
        self,
        file_path: Union[str, Path],
        law_id: Optional[str] = None,
        law_name: Optional[str] = None,
    ) -> ParseResult:
        """Synchronous wrapper for parse()."""
        return asyncio.run(self.parse(file_path, law_id, law_name))
    
    # =========================================================================
    # PDF Parsing
    # =========================================================================
    
    async def _parse_pdf(
        self,
        file_path: Path,
        law_id: Optional[str] = None,
        law_name: Optional[str] = None,
    ) -> ParseResult:
        """Parse a PDF document using LlamaParse."""
        result = ParseResult(success=False)
        
        logger.info(f"🔄 Parsing PDF: {file_path.name}")
        
        try:
            parser = self._get_llama_parser()
            
            # Parse with LlamaParse (run in thread pool)
            documents = await asyncio.to_thread(
                parser.load_data, str(file_path)
            )
            
            if not documents:
                result.errors.append("No content extracted from PDF")
                return result
            
            # Combine all document content
            raw_content = "\n\n".join(doc.text for doc in documents)
            result.raw_content = raw_content
            
            # Extract metadata
            law_id = law_id or self._extract_law_id(file_path.name, raw_content)
            law_name = law_name or self._extract_law_name(raw_content)
            doc_kind = self._detect_doc_kind(law_id, raw_content)
            
            # Extract tables
            result.tables = self._extract_tables(raw_content)
            
            # Create chunks using LlamaIndex nodes
            chunks = self._create_chunks(
                raw_content, 
                file_path.name, 
                law_id, 
                law_name, 
                doc_kind
            )
            result.chunks = chunks
            
            # Statistics
            result.statistics = {
                "total_tokens": count_tokens(raw_content, self._tokenizer),
                "total_chunks": len(chunks),
                "tables": len(result.tables),
                "chapters": len(re.findall(PATTERNS["chapter"], raw_content)),
                "articles": len(re.findall(PATTERNS["article"], raw_content)),
            }
            
            # Metadata
            result.metadata = {
                "doc_kind": doc_kind,
                "document_number": law_id,
                "title": law_name,
                "source_file": file_path.name,
                "parser": "llamaindex",
            }
            
            result.success = True
            logger.info(f"✅ Parsed PDF: {len(chunks)} chunks, {result.statistics['articles']} articles")
            
        except Exception as e:
            logger.error(f"❌ PDF parsing failed: {e}")
            
            # Fallback to PyPDF2
            if self.config.use_fallback_parser:
                logger.warning("⚠️ Falling back to PyPDF2")
                return await self._fallback_parse_pdf(file_path, law_id, law_name)
            
            result.errors.append(str(e))
        
        return result
    
    async def _fallback_parse_pdf(
        self,
        file_path: Path,
        law_id: Optional[str] = None,
        law_name: Optional[str] = None,
    ) -> ParseResult:
        """Fallback PDF parsing using PyPDF2."""
        result = ParseResult(success=False)
        
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(str(file_path))
            pages_text = []
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text() or ""
                pages_text.append(text)
            
            raw_content = "\n\n".join(pages_text)
            result.raw_content = raw_content
            result.warnings.append("Using PyPDF2 fallback - tables may be missing")
            
            # Extract metadata
            law_id = law_id or self._extract_law_id(file_path.name, raw_content)
            law_name = law_name or self._extract_law_name(raw_content)
            doc_kind = self._detect_doc_kind(law_id, raw_content)
            
            # Create chunks
            chunks = self._create_chunks(
                raw_content,
                file_path.name,
                law_id,
                law_name,
                doc_kind
            )
            result.chunks = chunks
            
            result.statistics = {
                "total_tokens": count_tokens(raw_content, self._tokenizer),
                "total_chunks": len(chunks),
                "pages": len(reader.pages),
            }
            
            result.metadata = {
                "doc_kind": doc_kind,
                "document_number": law_id,
                "title": law_name,
                "source_file": file_path.name,
                "parser": "pypdf2_fallback",
            }
            
            result.success = True
            
        except Exception as e:
            result.errors.append(f"PyPDF2 fallback failed: {e}")
        
        return result
    
    # =========================================================================
    # DOCX Parsing
    # =========================================================================
    
    async def _parse_docx(
        self,
        file_path: Path,
        law_id: Optional[str] = None,
        law_name: Optional[str] = None,
    ) -> ParseResult:
        """Parse a DOCX document."""
        result = ParseResult(success=False)
        
        logger.info(f"🔄 Parsing DOCX: {file_path.name}")
        
        try:
            # Try LlamaParse first for DOCX
            if self.config.llama_cloud_api_key:
                return await self._parse_pdf(file_path, law_id, law_name)
        except Exception as e:
            logger.warning(f"LlamaParse failed for DOCX: {e}")
        
        # Fallback to python-docx
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            paragraphs = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            for table in doc.tables:
                table_text = self._extract_docx_table(table)
                if table_text:
                    paragraphs.append(table_text)
            
            raw_content = "\n\n".join(paragraphs)
            
            # Normalize and clean
            raw_content = self._normalize_text(raw_content)
            result.raw_content = raw_content
            
            # Extract metadata
            law_id = law_id or self._extract_law_id(file_path.name, raw_content)
            law_name = law_name or self._extract_law_name(raw_content)
            doc_kind = self._detect_doc_kind(law_id, raw_content)
            
            # Create chunks
            chunks = self._create_chunks(
                raw_content,
                file_path.name,
                law_id,
                law_name,
                doc_kind
            )
            result.chunks = chunks
            
            # Get statistics from chunking (if available)
            stats = getattr(self, '_last_stats', {})
            
            result.statistics = {
                "total_tokens": count_tokens(raw_content, self._tokenizer),
                "total_chunks": len(chunks),
                "chapters": stats.get("chapters", len(re.findall(PATTERNS["chapter"], raw_content))),
                "sections": stats.get("sections", 0),
                "articles": stats.get("articles", len(re.findall(PATTERNS["article"], raw_content))),
                "clauses": stats.get("clauses", 0),
                "points": stats.get("points", 0),
            }
            
            result.metadata = {
                "doc_kind": doc_kind,
                "document_number": law_id,
                "title": law_name,
                "source_file": file_path.name,
                "parser": "python-docx",
            }
            
            result.success = True
            logger.info(f"✅ Parsed DOCX: {len(chunks)} chunks")
            
        except Exception as e:
            logger.error(f"❌ DOCX parsing failed: {e}")
            result.errors.append(str(e))
        
        return result
    
    def _extract_docx_table(self, table) -> str:
        """Convert a DOCX table to Markdown format."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        
        if len(rows) > 1:
            # Add header separator
            header = rows[0]
            separator = "| " + " | ".join(["---"] * len(table.rows[0].cells)) + " |"
            body = "\n".join(rows[1:])
            return f"{header}\n{separator}\n{body}"
        
        return "\n".join(rows)
    
    # =========================================================================
    # Chunking - Improved hierarchical chunking like legacy parser
    # =========================================================================
    
    def _create_chunks(
        self,
        content: str,
        source_file: str,
        law_id: str,
        law_name: str,
        doc_kind: str,
    ) -> List[LegalChunk]:
        """
        Create chunks from document content.
        
        Uses hierarchical chunking similar to legacy parser:
        1. Build document tree (Chapter > Section > Article > Clause > Point)
        2. Generate chunks with proper metadata and relationships
        3. Split by token threshold when needed
        """
        # Use hierarchical tree-based chunking (like legacy parser)
        return self._create_hierarchical_chunks(
            content, source_file, law_id, law_name, doc_kind
        )
    
    def _create_hierarchical_chunks(
        self,
        content: str,
        source_file: str,
        law_id: str,
        law_name: str,
        doc_kind: str,
    ) -> List[LegalChunk]:
        """
        Create chunks using hierarchical tree structure.
        This replicates the logic from legacy VietnamLegalDocxParser.
        """
        chunks = []
        lines = content.split('\n')
        
        # Current context tracking
        current_chapter = None
        current_section = None
        current_article = None
        current_clause = None
        
        # Statistics tracking
        stats = {
            "chapters": 0,
            "sections": 0,
            "articles": 0,
            "clauses": 0,
            "points": 0,
        }
        
        # Accumulated content for current context
        article_buffer = []
        article_clauses = []  # Store clause chunks for potential merging
        clause_buffer = []
        
        def flush_clause():
            """Flush accumulated clause content - store for article merging."""
            nonlocal clause_buffer, current_clause, article_clauses
            if clause_buffer and current_clause:
                clause_content = ' '.join(clause_buffer).strip()
                if clause_content:
                    chunk = self._create_chunk_from_context(
                        content=clause_content,
                        source_file=source_file,
                        law_id=law_id,
                        law_name=law_name,
                        doc_kind=doc_kind,
                        chapter=current_chapter,
                        section=current_section,
                        article=current_article,
                        clause=current_clause,
                        point=None,
                        is_definition=current_article and current_article.get('is_definition', False),
                    )
                    article_clauses.append(chunk)
            clause_buffer = []
        
        def flush_article():
            """
            Flush accumulated article content.
            If total tokens <= threshold, merge into single chunk.
            Otherwise, add individual clause chunks.
            """
            nonlocal article_buffer, article_clauses, current_article
            flush_clause()  # Flush any pending clause first
            
            if current_article:
                # Calculate total article tokens
                total_content = ' '.join(article_buffer).strip()
                total_tokens = count_tokens(total_content, self._tokenizer)
                
                if total_tokens <= self.config.chunk_size:
                    # Article fits in one chunk - create single article chunk
                    if total_content:
                        chunk = self._create_chunk_from_context(
                            content=total_content,
                            source_file=source_file,
                            law_id=law_id,
                            law_name=law_name,
                            doc_kind=doc_kind,
                            chapter=current_chapter,
                            section=current_section,
                            article=current_article,
                            clause=None,
                            point=None,
                            is_definition=current_article.get('is_definition', False),
                        )
                        chunks.append(chunk)
                else:
                    # Article too large - add individual clause chunks
                    chunks.extend(article_clauses)
            
            article_buffer = []
            article_clauses = []
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
            
            # Check for Chapter (Chương)
            chapter_match = PATTERNS["chapter"].match(line)
            if chapter_match:
                flush_article()
                current_chapter = {
                    "id": chapter_match.group(1).upper(),
                    "title": chapter_match.group(2).strip(),
                }
                current_section = None
                current_article = None
                current_clause = None
                stats["chapters"] += 1
                i += 1
                continue
            
            # Check for Section (Mục)
            section_match = PATTERNS["section"].match(line)
            if section_match:
                flush_article()
                current_section = {
                    "id": section_match.group(1),
                    "title": section_match.group(2).strip(),
                }
                current_article = None
                current_clause = None
                stats["sections"] += 1
                i += 1
                continue
            
            # Check for Article (Điều)
            article_match = PATTERNS["article"].match(line)
            if article_match:
                flush_article()
                article_id = article_match.group(1)
                article_title = article_match.group(2).strip()
                
                # Check if this is a definition article
                is_definition = False
                look_ahead = ' '.join(lines[i:i+5])
                for pattern in DEFINITION_PATTERNS:
                    if pattern.search(look_ahead):
                        is_definition = True
                        break
                
                current_article = {
                    "id": article_id,
                    "title": article_title,
                    "is_definition": is_definition,
                }
                current_clause = None
                article_buffer = [line]  # Start with article header
                stats["articles"] += 1
                i += 1
                continue
            
            # Check for Clause (Khoản) - only within article context
            if current_article:
                clause_match = PATTERNS["clause"].match(line)
                if clause_match:
                    clause_num = int(clause_match.group(1))
                    # Validate: clause numbers should be reasonable (1-99)
                    if clause_num < 100:
                        flush_clause()
                        current_clause = {
                            "id": clause_match.group(1),
                            "content": clause_match.group(2).strip(),
                        }
                        clause_buffer = [line]
                        article_buffer.append(line)  # Also track in article buffer
                        stats["clauses"] += 1
                        i += 1
                        continue
            
            # Check for Point (Điểm) - only within clause context
            # Points are accumulated into clause buffer, only split if too large
            if current_clause:
                point_match = PATTERNS["point"].match(line)
                if point_match:
                    stats["points"] += 1
                    article_buffer.append(line)  # Also track in article buffer
                    
                    # Check if current clause is getting too large
                    current_tokens = count_tokens(' '.join(clause_buffer), self._tokenizer)
                    line_tokens = count_tokens(line, self._tokenizer)
                    
                    if current_tokens + line_tokens > self.config.chunk_size:
                        # Flush current clause and start new chunk with this point
                        flush_clause()
                        
                        # Create a new clause context for overflowing points
                        point_id = point_match.group(1).lower()
                        chunk = self._create_chunk_from_context(
                            content=line,
                            source_file=source_file,
                            law_id=law_id,
                            law_name=law_name,
                            doc_kind=doc_kind,
                            chapter=current_chapter,
                            section=current_section,
                            article=current_article,
                            clause=current_clause,
                            point={"id": point_id, "content": point_match.group(2).strip()},
                            is_definition=current_article and current_article.get('is_definition', False),
                        )
                        article_clauses.append(chunk)  # Store for potential article merge
                        clause_buffer = []
                    else:
                        # Add point to current clause buffer
                        clause_buffer.append(line)
                    
                    i += 1
                    continue
            
            # Accumulate content to appropriate buffer
            if current_clause:
                clause_buffer.append(line)
                article_buffer.append(line)  # Also track in article buffer
            elif current_article:
                article_buffer.append(line)
            
            i += 1
        
        # Flush any remaining content
        flush_article()
        
        # Store statistics for later
        self._last_stats = stats
        
        logger.info(f"Created {len(chunks)} chunks: {stats['articles']} articles, {stats['clauses']} clauses, {stats['points']} points")
        
        return chunks
    
    def _create_chunk_from_context(
        self,
        content: str,
        source_file: str,
        law_id: str,
        law_name: str,
        doc_kind: str,
        chapter: Optional[Dict],
        section: Optional[Dict],
        article: Optional[Dict],
        clause: Optional[Dict],
        point: Optional[Dict],
        is_definition: bool = False,
    ) -> LegalChunk:
        """Create a LegalChunk from the current parsing context."""
        
        # Build structure dict
        structure = {}
        
        if chapter:
            structure["chapter_id"] = chapter["id"]
            structure["chapter_title"] = chapter.get("title", "")
            structure["chapter"] = f"Chương {chapter['id']}"
            if chapter.get("title"):
                structure["chapter"] += f" {chapter['title']}"
        
        if section:
            structure["section_id"] = section["id"]
            structure["section_title"] = section.get("title", "")
            structure["section"] = f"Mục {section['id']}"
            if section.get("title"):
                structure["section"] += f" {section['title']}"
        
        if article:
            structure["article_id"] = f"Điều {article['id']}"
            structure["article_number"] = article["id"]
            structure["article_title"] = article.get("title", "")
        
        if clause:
            structure["clause_no"] = clause["id"]
        
        if point:
            structure["point_no"] = point["id"]
        
        # Build chunk_id
        chunk_id_parts = [f"LAW={law_id}"]
        if chapter:
            chunk_id_parts.append(f"CHUONG={chapter['id']}")
        if section:
            chunk_id_parts.append(f"MUC={section['id']}")
        if article:
            chunk_id_parts.append(f"DIEU={article['id']}")
        if clause:
            if is_definition:
                chunk_id_parts.append(f"DEFINITION_ITEM={clause['id']}")
            else:
                chunk_id_parts.append(f"KHOAN={clause['id']}")
        if point:
            chunk_id_parts.append(f"DIEM={point['id']}")
        
        chunk_id = ":".join(chunk_id_parts)
        
        # Build embedding prefix
        embedding_prefix = self._build_embedding_prefix(law_id, structure, doc_kind)
        
        # Build metadata
        metadata = {
            "doc_kind": doc_kind,
            "document_number": law_id,
            "law_id": law_id,
            "law_name": law_name,
            "source_file": source_file,
            "is_definition": is_definition,
            **{k: v for k, v in structure.items() if v is not None},
        }
        
        # Calculate tokens
        tokens = count_tokens(content, self._tokenizer)
        
        return LegalChunk(
            chunk_id=chunk_id,
            content=content.strip(),
            embedding_prefix=embedding_prefix,
            metadata=metadata,
            tokens=tokens,
        )
    
    # =========================================================================
    # Helper Methods
    # =========================================================================
    
    def _normalize_text(self, text: str) -> str:
        """Normalize Vietnamese text."""
        # Unicode normalization (NFC for Vietnamese)
        text = unicodedata.normalize("NFC", text)
        
        # Remove boilerplate
        boilerplate_patterns = [
            r"^\s*QUỐC\s*HỘI\s*$",
            r"^\s*CỘNG\s*HÒA\s*XÃ\s*HỘI\s*CHỦ\s*NGHĨA\s*VIỆT\s*NAM\s*$",
            r"^\s*Độc\s*lập.*Tự\s*do.*Hạnh\s*phúc\s*$",
            r"^\s*-{3,}\s*$",
            r"^\s*_{3,}\s*$",
            r"^\s*Trang\s*\d+\s*$",
            r"^\s*\d+\s*/\s*\d+\s*$",
        ]
        
        for pattern in boilerplate_patterns:
            text = re.sub(pattern, "", text, flags=re.IGNORECASE | re.MULTILINE)
        
        # Normalize whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        return text.strip()
    
    def _extract_law_id(self, filename: str, content: str) -> str:
        """Extract law ID from filename or content."""
        # Try filename first
        match = LAW_ID_PATTERN.search(filename)
        if match:
            parts = [g for g in match.groups() if g]
            return "/".join(parts)
        
        # Try content
        for line in content.split('\n')[:30]:
            match = LAW_ID_PATTERN.search(line)
            if match:
                parts = [g for g in match.groups() if g]
                return "/".join(parts)
        
        return "UNKNOWN"
    
    def _extract_law_name(self, content: str) -> str:
        """Extract law name from content."""
        patterns = [
            (r"^\s*LUẬT\s+(.+)$", "Luật"),
            (r"^\s*NGHỊ\s*ĐỊNH\s+(.+)$", "Nghị định"),
            (r"^\s*THÔNG\s*TƯ\s+(.+)$", "Thông tư"),
        ]
        
        for line in content.split('\n')[:50]:
            line = line.strip()
            for pattern, prefix in patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    return f"{prefix} {match.group(1).strip()}"
        
        return "Unknown Document"
    
    def _detect_doc_kind(self, law_id: str, content: str) -> str:
        """Detect document type (LAW, DECREE, CIRCULAR)."""
        law_id_upper = (law_id or "").upper()
        
        # Check law_id patterns
        if "QH" in law_id_upper:
            return "LAW"
        if "NĐ-CP" in law_id_upper or "ND-CP" in law_id_upper:
            return "DECREE"
        if "TT-" in law_id_upper:
            return "CIRCULAR"
        
        # Check content patterns
        for doc_type, patterns in DOC_TYPE_PATTERNS.items():
            for pattern in patterns:
                if pattern.search(content[:2000]):
                    return doc_type
        
        return "LAW"  # Default
    
    def _extract_structure_from_text(self, text: str) -> Dict[str, Any]:
        """Extract legal structure metadata from text."""
        structure = {}
        
        # Check for chapter
        chapter_match = PATTERNS["chapter"].search(text)
        if chapter_match:
            structure["chapter_id"] = chapter_match.group(1)
            structure["chapter_title"] = chapter_match.group(2).strip()
        
        # Check for section
        section_match = PATTERNS["section"].search(text)
        if section_match:
            structure["section_id"] = section_match.group(1)
            structure["section_title"] = section_match.group(2).strip()
        
        # Check for article
        article_match = PATTERNS["article"].search(text)
        if article_match:
            structure["article_id"] = f"Điều {article_match.group(1)}"
            structure["article_number"] = article_match.group(1)
            structure["article_title"] = article_match.group(2).strip()
        
        return structure
    
    def _build_embedding_prefix(
        self,
        law_id: str,
        structure: Dict[str, Any],
        doc_kind: str,
    ) -> str:
        """Build embedding prefix for the chunk."""
        parts = []
        
        parts.append(f"DOC={law_id}")
        parts.append(f"TYPE={doc_kind}")
        
        if structure.get("chapter_id"):
            parts.append(f"CHUONG={structure['chapter_id']}")
        if structure.get("section_id"):
            parts.append(f"MUC={structure['section_id']}")
        if structure.get("article_number"):
            parts.append(f"DIEU={structure['article_number']}")
        if structure.get("clause_no"):
            parts.append(f"KHOAN={structure['clause_no']}")
        if structure.get("point_no"):
            parts.append(f"DIEM={structure['point_no']}")
        
        return " | ".join(parts)
    
    def _generate_chunk_id(
        self,
        law_id: str,
        structure: Dict[str, Any],
        idx: int,
    ) -> str:
        """Generate a unique chunk ID."""
        parts = [f"LAW={law_id}"]
        
        if structure.get("chapter_id"):
            parts.append(f"CHUONG={structure['chapter_id']}")
        if structure.get("section_id"):
            parts.append(f"MUC={structure['section_id']}")
        if structure.get("article_number"):
            parts.append(f"DIEU={structure['article_number']}")
        if structure.get("clause_no"):
            parts.append(f"KHOAN={structure['clause_no']}")
        if structure.get("point_no"):
            parts.append(f"DIEM={structure['point_no']}")
        
        base_id = ":".join(parts)
        
        # Add hash for uniqueness
        content_hash = hashlib.md5(f"{base_id}:{idx}".encode()).hexdigest()[:8]
        
        return f"{base_id}:{content_hash}"
    
    def _extract_tables(self, content: str) -> List[Dict[str, Any]]:
        """Extract Markdown tables from content."""
        tables = []
        
        # Markdown table pattern
        table_pattern = r'(\|[^\n]+\|\n)(\|[-:| ]+\|\n)((?:\|[^\n]+\|\n)*)'
        
        for i, match in enumerate(re.finditer(table_pattern, content)):
            header = match.group(1).strip()
            body = match.group(3).strip()
            
            # Parse columns
            columns = [col.strip() for col in header.split('|')[1:-1]]
            
            # Parse rows
            rows = []
            for row_line in body.split('\n'):
                if row_line.strip():
                    cells = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    rows.append(cells)
            
            tables.append({
                "id": f"table_{i+1}",
                "columns": columns,
                "rows": rows,
                "num_rows": len(rows),
                "num_cols": len(columns),
                "markdown": match.group(0),
            })
        
        return tables


# =============================================================================
# Convenience Functions
# =============================================================================

def parse_legal_document(
    file_path: Union[str, Path],
    law_id: Optional[str] = None,
    law_name: Optional[str] = None,
) -> ParseResult:
    """
    Convenience function to parse a Vietnamese legal document.
    
    Args:
        file_path: Path to the document (PDF or DOCX)
        law_id: Override law ID
        law_name: Override law name
        
    Returns:
        ParseResult with chunks and metadata
    """
    parser = LlamaIndexLegalParser.from_env()
    return parser.parse_sync(file_path, law_id, law_name)


async def parse_legal_document_async(
    file_path: Union[str, Path],
    law_id: Optional[str] = None,
    law_name: Optional[str] = None,
) -> ParseResult:
    """Async version of parse_legal_document."""
    parser = LlamaIndexLegalParser.from_env()
    return await parser.parse(file_path, law_id, law_name)


# =============================================================================
# CLI
# =============================================================================

if __name__ == "__main__":
    import sys
    
    logging.basicConfig(level=logging.INFO)
    
    if len(sys.argv) < 2:
        print("Usage: python llamaindex_legal_parser.py <file.pdf|file.docx>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    result = parse_legal_document(file_path)
    
    if result.success:
        print(f"✓ Parsed successfully!")
        print(f"  Total chunks: {result.statistics.get('total_chunks', 0)}")
        print(f"  Total tokens: {result.statistics.get('total_tokens', 0)}")
        print(f"  Articles: {result.statistics.get('articles', 0)}")
        
        print("\nFirst 3 chunks:")
        for chunk in result.chunks[:3]:
            print(f"\n--- {chunk.chunk_id} ---")
            print(f"Prefix: {chunk.embedding_prefix}")
            print(f"Tokens: {chunk.tokens}")
            print(f"Content: {chunk.content[:200]}...")
    else:
        print(f"✗ Parse failed!")
        for error in result.errors:
            print(f"  Error: {error}")
