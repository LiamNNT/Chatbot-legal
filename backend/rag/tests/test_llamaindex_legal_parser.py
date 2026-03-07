# tests/test_llamaindex_legal_parser.py
"""
Unit tests for LlamaIndexLegalParser.

Tests cover:
- Domain models (LegalNodeType, LegalNode, LegalChunk, ParseResult)
- Hierarchical chunking logic
- Tree building for Knowledge Graph
- Metadata extraction (law_id, law_name, doc_kind)
- Embedding prefix generation
- Vietnamese legal structure detection (Chương, Mục, Điều, Khoản, Điểm)
- Token counting
- DOCX parsing (python-docx fallback)
- Edge cases
"""

from __future__ import annotations

import asyncio
import os
import tempfile
from pathlib import Path
from typing import List, Optional
from unittest.mock import AsyncMock, MagicMock, patch

import pytest

from app.ingest.loaders.llamaindex_legal_parser import (
    DEFINITION_PATTERNS,
    PATTERNS,
    LegalChunk,
    LegalNode,
    LegalNodeType,
    LlamaIndexLegalParser,
    ParseResult,
    ParserConfig,
    count_tokens,
)


# =============================================================================
# Fixtures
# =============================================================================


@pytest.fixture
def parser_config() -> ParserConfig:
    """Default parser config for testing (no LlamaParse API key)."""
    return ParserConfig(
        llama_cloud_api_key=None,
        chunk_size=800,
        chunk_overlap=100,
        use_fallback_parser=True,
    )


@pytest.fixture
def parser(parser_config: ParserConfig) -> LlamaIndexLegalParser:
    return LlamaIndexLegalParser(parser_config)


SAMPLE_LAW_TEXT = """\
QUỐC HỘI
CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM
Độc lập - Tự do - Hạnh phúc

LUẬT GIÁO DỤC ĐẠI HỌC

Chương I QUY ĐỊNH CHUNG

Điều 1. Phạm vi điều chỉnh
Luật này quy định về tổ chức, nhiệm vụ, quyền hạn của cơ sở giáo dục đại học.

Điều 2. Đối tượng áp dụng
1. Cơ sở giáo dục đại học.
2. Tổ chức, cá nhân có liên quan đến giáo dục đại học.

Điều 3. Giải thích từ ngữ
Trong Luật này, các từ ngữ dưới đây được hiểu như sau:
1. Giáo dục đại học là bậc giáo dục sau trung học phổ thông.
2. Cơ sở giáo dục đại học bao gồm trường đại học và học viện.

Chương II TỔ CHỨC CƠ SỞ GIÁO DỤC ĐẠI HỌC

Mục 1 Loại hình cơ sở giáo dục đại học

Điều 4. Cơ sở giáo dục đại học
1. Đại học quốc gia, đại học vùng.
2. Trường đại học, học viện.
a) Trường đại học công lập.
b) Trường đại học tư thục.

Điều 5. Quyền tự chủ đại học
1. Tự chủ trong học thuật.
2. Tự chủ trong tổ chức và nhân sự.
3. Tự chủ trong tài chính và tài sản.
"""


# =============================================================================
# Domain Model Tests
# =============================================================================


class TestLegalNodeType:
    """Tests for LegalNodeType enum."""

    def test_enum_values(self):
        assert LegalNodeType.LAW == "LAW"
        assert LegalNodeType.CHAPTER == "CHUONG"
        assert LegalNodeType.SECTION == "MUC"
        assert LegalNodeType.ARTICLE == "DIEU"
        assert LegalNodeType.CLAUSE == "KHOAN"
        assert LegalNodeType.POINT == "DIEM"
        assert LegalNodeType.DEFINITION_ITEM == "DEFINITION_ITEM"

    def test_is_string_enum(self):
        """LegalNodeType inherits from str, so .value can be used as string."""
        assert LegalNodeType.LAW.value == "LAW"
        assert f"TYPE={LegalNodeType.LAW.value}" == "TYPE=LAW"


class TestLegalNode:
    """Tests for LegalNode dataclass."""

    def test_basic_creation(self):
        node = LegalNode(
            node_type=LegalNodeType.ARTICLE,
            identifier="1",
            title="Phạm vi điều chỉnh",
            content="Luật này quy định...",
        )
        assert node.node_type == LegalNodeType.ARTICLE
        assert node.identifier == "1"
        assert node.title == "Phạm vi điều chỉnh"
        assert node.children == []
        assert node.parent is None

    def test_get_full_id(self):
        root = LegalNode(node_type=LegalNodeType.LAW, identifier="24/2018/QH14")
        chapter = LegalNode(node_type=LegalNodeType.CHAPTER, identifier="I", parent=root)
        article = LegalNode(node_type=LegalNodeType.ARTICLE, identifier="1", parent=chapter)

        assert article.get_full_id() == "LAW=24/2018/QH14:CHUONG=I:DIEU=1"
        assert chapter.get_full_id() == "LAW=24/2018/QH14:CHUONG=I"
        assert root.get_full_id() == "LAW=24/2018/QH14"

    def test_get_lineage(self):
        root = LegalNode(node_type=LegalNodeType.LAW, identifier="X")
        chapter = LegalNode(node_type=LegalNodeType.CHAPTER, identifier="I", parent=root)
        article = LegalNode(node_type=LegalNodeType.ARTICLE, identifier="1", parent=chapter)

        assert article.get_lineage() == ["LAW", "CHUONG", "DIEU"]

    def test_get_ancestors_include_self(self):
        root = LegalNode(node_type=LegalNodeType.LAW, identifier="24/2018/QH14", title="Luật GDĐH")
        chapter = LegalNode(node_type=LegalNodeType.CHAPTER, identifier="II", title="Tổ chức", parent=root)
        article = LegalNode(node_type=LegalNodeType.ARTICLE, identifier="4", title="Cơ sở GDĐH", parent=chapter)

        ancestors = article.get_ancestors(include_self=True)
        assert ancestors["law_id"] == "24/2018/QH14"
        assert ancestors["law_name"] == "Luật GDĐH"
        assert ancestors["chapter_id"] == "II"
        assert ancestors["article_id"] == "Điều 4"
        assert ancestors["article_number"] == "4"

    def test_get_ancestors_exclude_self(self):
        root = LegalNode(node_type=LegalNodeType.LAW, identifier="X", title="Test")
        article = LegalNode(node_type=LegalNodeType.ARTICLE, identifier="1", parent=root)

        ancestors = article.get_ancestors(include_self=False)
        assert "article_id" not in ancestors
        assert ancestors["law_id"] == "X"

    def test_tree_relationships(self):
        root = LegalNode(node_type=LegalNodeType.LAW, identifier="X")
        child1 = LegalNode(node_type=LegalNodeType.CHAPTER, identifier="I", parent=root)
        child2 = LegalNode(node_type=LegalNodeType.CHAPTER, identifier="II", parent=root)
        root.children = [child1, child2]
        child1.next_sibling = child2
        child2.prev_sibling = child1

        assert root.children[0].identifier == "I"
        assert root.children[1].identifier == "II"
        assert child1.next_sibling is child2
        assert child2.prev_sibling is child1

    def test_definition_article(self):
        node = LegalNode(
            node_type=LegalNodeType.ARTICLE,
            identifier="3",
            title="Giải thích từ ngữ",
            is_definition_article=True,
        )
        assert node.is_definition_article is True


class TestLegalChunk:
    """Tests for LegalChunk dataclass."""

    def test_creation(self):
        chunk = LegalChunk(
            chunk_id="LAW=X:DIEU=1",
            content="Luật này quy định...",
            embedding_prefix="DOC=X | TYPE=LAW | DIEU=1",
            metadata={"law_id": "X", "article_id": "Điều 1"},
            tokens=15,
        )
        assert chunk.chunk_id == "LAW=X:DIEU=1"
        assert chunk.tokens == 15

    def test_to_dict(self):
        chunk = LegalChunk(
            chunk_id="id1",
            content="test content",
            embedding_prefix="prefix",
            metadata={"key": "val"},
            tokens=5,
        )
        d = chunk.to_dict()
        assert d["chunk_id"] == "id1"
        assert d["content"] == "test content"
        assert d["tokens"] == 5


class TestParseResult:
    """Tests for ParseResult dataclass."""

    def test_default_values(self):
        result = ParseResult(success=True)
        assert result.success is True
        assert result.chunks == []
        assert result.tree is None
        assert result.errors == []

    def test_with_tree(self):
        tree = LegalNode(node_type=LegalNodeType.LAW, identifier="X")
        result = ParseResult(success=True, tree=tree)
        assert result.tree is not None
        assert result.tree.identifier == "X"


# =============================================================================
# Pattern Detection Tests
# =============================================================================


class TestPatterns:
    """Tests for Vietnamese legal structure regex patterns."""

    def test_chapter_pattern(self):
        assert PATTERNS["chapter"].match("Chương I QUY ĐỊNH CHUNG")
        assert PATTERNS["chapter"].match("CHƯƠNG II TỔ CHỨC")
        assert PATTERNS["chapter"].match("Chương III")
        assert not PATTERNS["chapter"].match("Điều 1. Phạm vi")

    def test_section_pattern(self):
        assert PATTERNS["section"].match("Mục 1 Loại hình cơ sở")
        assert PATTERNS["section"].match("MỤC 2 Tổ chức")
        assert not PATTERNS["section"].match("Điều 1.")

    def test_article_pattern(self):
        m = PATTERNS["article"].match("Điều 1. Phạm vi điều chỉnh")
        assert m is not None
        assert m.group(1) == "1"
        assert m.group(2) == "Phạm vi điều chỉnh"

        m2 = PATTERNS["article"].match("Điều 10a. Bổ sung")
        assert m2 is not None
        assert m2.group(1) == "10a"

    def test_clause_pattern(self):
        m = PATTERNS["clause"].match("1. Cơ sở giáo dục đại học.")
        assert m is not None
        assert m.group(1) == "1"

        assert not PATTERNS["clause"].match("100. Số quá lớn")

    def test_point_pattern(self):
        m = PATTERNS["point"].match("a) Trường đại học công lập.")
        assert m is not None
        assert m.group(1) == "a"

        m2 = PATTERNS["point"].match("đ) Trường hợp đặc biệt.")
        assert m2 is not None

    def test_definition_patterns(self):
        text = "Trong Luật này, các từ ngữ dưới đây được hiểu như sau"
        found = any(p.search(text) for p in DEFINITION_PATTERNS)
        assert found

        text2 = "Giải thích từ ngữ"
        found2 = any(p.search(text2) for p in DEFINITION_PATTERNS)
        assert found2


# =============================================================================
# Token Counting Tests
# =============================================================================


class TestTokenCounting:
    """Tests for token counting utility."""

    def test_empty_text(self):
        assert count_tokens("") == 0
        assert count_tokens(None) == 0

    def test_nonempty_text(self):
        tokens = count_tokens("Luật này quy định về tổ chức")
        assert tokens > 0

    def test_longer_text_has_more_tokens(self):
        short = count_tokens("Điều 1")
        long = count_tokens("Điều 1. Phạm vi điều chỉnh của Luật giáo dục đại học")
        assert long > short


# =============================================================================
# Hierarchical Chunking Tests
# =============================================================================


class TestHierarchicalChunking:
    """Tests for the hierarchical chunking and tree building."""

    def test_basic_chunking(self, parser: LlamaIndexLegalParser):
        chunks, tree = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="24/2018/QH14",
            law_name="Luật GDĐH",
            doc_kind="LAW",
        )

        assert len(chunks) > 0
        assert tree is not None
        assert tree.node_type == LegalNodeType.LAW
        assert tree.identifier == "24/2018/QH14"

    def test_tree_has_chapters(self, parser: LlamaIndexLegalParser):
        _, tree = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )

        chapter_nodes = [c for c in tree.children if c.node_type == LegalNodeType.CHAPTER]
        assert len(chapter_nodes) == 2
        assert chapter_nodes[0].identifier == "I"
        assert chapter_nodes[1].identifier == "II"

    def test_tree_has_articles(self, parser: LlamaIndexLegalParser):
        _, tree = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )

        # Chapter I should have articles 1, 2, 3
        chap1 = tree.children[0]
        article_nodes = [c for c in chap1.children if c.node_type == LegalNodeType.ARTICLE]
        assert len(article_nodes) == 3
        assert article_nodes[0].identifier == "1"
        assert article_nodes[1].identifier == "2"
        assert article_nodes[2].identifier == "3"

    def test_tree_has_sections(self, parser: LlamaIndexLegalParser):
        _, tree = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )

        # Chapter II should have Section 1
        chap2 = tree.children[1]
        section_nodes = [c for c in chap2.children if c.node_type == LegalNodeType.SECTION]
        assert len(section_nodes) == 1
        assert section_nodes[0].identifier == "1"

    def test_tree_has_clauses(self, parser: LlamaIndexLegalParser):
        _, tree = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )

        # Điều 2 should have 2 clauses
        chap1 = tree.children[0]
        article2 = [c for c in chap1.children if c.identifier == "2"][0]
        clause_nodes = [c for c in article2.children if c.node_type == LegalNodeType.CLAUSE]
        assert len(clause_nodes) == 2

    def test_tree_has_points(self, parser: LlamaIndexLegalParser):
        _, tree = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )

        # Điều 4, Khoản 2 should have points a, b
        chap2 = tree.children[1]
        section1 = [c for c in chap2.children if c.node_type == LegalNodeType.SECTION][0]
        article4 = [c for c in section1.children if c.identifier == "4"][0]
        clause2 = [c for c in article4.children if c.identifier == "2"][0]
        point_nodes = [c for c in clause2.children if c.node_type == LegalNodeType.POINT]
        assert len(point_nodes) == 2
        assert point_nodes[0].identifier == "a"
        assert point_nodes[1].identifier == "b"

    def test_definition_article_detected(self, parser: LlamaIndexLegalParser):
        _, tree = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )

        chap1 = tree.children[0]
        article3 = [c for c in chap1.children if c.identifier == "3"][0]
        assert article3.is_definition_article is True

    def test_chunk_metadata(self, parser: LlamaIndexLegalParser):
        chunks, _ = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="24/2018/QH14",
            law_name="Luật GDĐH",
            doc_kind="LAW",
        )

        # All chunks should have law_id in metadata
        for chunk in chunks:
            assert chunk.metadata["law_id"] == "24/2018/QH14"
            assert chunk.metadata["doc_kind"] == "LAW"

    def test_chunk_embedding_prefix(self, parser: LlamaIndexLegalParser):
        chunks, _ = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="24/2018/QH14",
            law_name="Test",
            doc_kind="LAW",
        )

        for chunk in chunks:
            assert "DOC=24/2018/QH14" in chunk.embedding_prefix
            assert "TYPE=LAW" in chunk.embedding_prefix

    def test_chunk_ids_unique(self, parser: LlamaIndexLegalParser):
        chunks, _ = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )

        ids = [c.chunk_id for c in chunks]
        assert len(ids) == len(set(ids)), f"Duplicate chunk IDs found: {ids}"

    def test_sibling_links_in_tree(self, parser: LlamaIndexLegalParser):
        _, tree = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )

        # Chapters should be linked as siblings
        chap1, chap2 = tree.children[0], tree.children[1]
        assert chap1.next_sibling is chap2
        assert chap2.prev_sibling is chap1
        assert chap1.prev_sibling is None
        assert chap2.next_sibling is None

    def test_parent_links_in_tree(self, parser: LlamaIndexLegalParser):
        _, tree = parser._create_hierarchical_chunks(
            content=SAMPLE_LAW_TEXT,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )

        for child in tree.children:
            assert child.parent is tree


# =============================================================================
# Metadata Extraction Tests
# =============================================================================


class TestMetadataExtraction:
    """Tests for law ID, name, and doc_kind extraction."""

    def test_extract_law_id_from_filename(self, parser: LlamaIndexLegalParser):
        law_id = parser._extract_law_id("Luat-24-2018-QH14.docx", "")
        assert "24" in law_id
        assert "2018" in law_id

    def test_extract_law_id_from_content(self, parser: LlamaIndexLegalParser):
        content = "Căn cứ Luật số 24/2018/QH14 ngày 1/1/2018"
        law_id = parser._extract_law_id("unknown.docx", content)
        assert "24" in law_id
        assert "2018" in law_id

    def test_extract_law_id_unknown(self, parser: LlamaIndexLegalParser):
        law_id = parser._extract_law_id("random_file.docx", "no law id here")
        assert law_id == "UNKNOWN"

    def test_extract_law_name(self, parser: LlamaIndexLegalParser):
        content = "LUẬT GIÁO DỤC ĐẠI HỌC\nĐiều 1."
        name = parser._extract_law_name(content)
        assert "Luật" in name
        assert "GIÁO DỤC ĐẠI HỌC" in name

    def test_detect_doc_kind_law(self, parser: LlamaIndexLegalParser):
        assert parser._detect_doc_kind("24/2018/QH14", "") == "LAW"

    def test_detect_doc_kind_decree(self, parser: LlamaIndexLegalParser):
        assert parser._detect_doc_kind("100/2020/NĐ-CP", "") == "DECREE"

    def test_detect_doc_kind_circular(self, parser: LlamaIndexLegalParser):
        assert parser._detect_doc_kind("10/2021/TT-BGDĐT", "") == "CIRCULAR"

    def test_detect_doc_kind_from_content(self, parser: LlamaIndexLegalParser):
        content = "NGHỊ ĐỊNH\nVề tổ chức giáo dục"
        assert parser._detect_doc_kind("unknown", content) == "DECREE"


# =============================================================================
# Embedding Prefix Tests
# =============================================================================


class TestEmbeddingPrefix:
    """Tests for embedding prefix generation."""

    def test_basic_prefix(self, parser: LlamaIndexLegalParser):
        prefix = parser._build_embedding_prefix(
            "24/2018/QH14",
            {"chapter_id": "I", "article_number": "1"},
            "LAW",
        )
        assert "DOC=24/2018/QH14" in prefix
        assert "CHUONG=I" in prefix
        assert "DIEU=1" in prefix

    def test_prefix_with_clause_and_point(self, parser: LlamaIndexLegalParser):
        prefix = parser._build_embedding_prefix(
            "X",
            {"clause_no": "2", "point_no": "a"},
            "LAW",
        )
        assert "KHOAN=2" in prefix
        assert "DIEM=a" in prefix


# =============================================================================
# Full Parse (Async) Tests
# =============================================================================


class TestParseAsync:
    """Tests for the main async parse() method."""

    @pytest.mark.asyncio
    async def test_parse_nonexistent_file(self, parser: LlamaIndexLegalParser):
        result = await parser.parse(Path("/nonexistent/file.docx"))
        assert result.success is False
        assert len(result.errors) > 0

    @pytest.mark.asyncio
    async def test_parse_unsupported_extension(self, parser: LlamaIndexLegalParser):
        with tempfile.NamedTemporaryFile(suffix=".txt", delete=False) as f:
            f.write(b"test")
            f.flush()
            result = await parser.parse(Path(f.name))
        os.unlink(f.name)
        assert result.success is False

    @pytest.mark.asyncio
    async def test_parse_docx_fallback(self, parser: LlamaIndexLegalParser):
        """Test DOCX parsing via python-docx fallback (no LlamaParse key)."""
        try:
            from docx import Document

            # Create a minimal DOCX file
            doc = Document()
            doc.add_paragraph("LUẬT GIÁO DỤC ĐẠI HỌC")
            doc.add_paragraph("Chương I QUY ĐỊNH CHUNG")
            doc.add_paragraph("Điều 1. Phạm vi điều chỉnh")
            doc.add_paragraph("Luật này quy định về tổ chức, nhiệm vụ.")
            doc.add_paragraph("Điều 2. Đối tượng áp dụng")
            doc.add_paragraph("1. Cơ sở giáo dục đại học.")
            doc.add_paragraph("2. Tổ chức, cá nhân có liên quan.")

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                doc.save(f.name)
                fpath = Path(f.name)

            result = await parser.parse(fpath)
            os.unlink(fpath)

            assert result.success is True
            assert len(result.chunks) > 0
            assert result.tree is not None
            assert result.tree.node_type == LegalNodeType.LAW
            assert result.metadata.get("parser") == "python-docx"

            # Check tree structure
            chapter_nodes = [
                c for c in result.tree.children if c.node_type == LegalNodeType.CHAPTER
            ]
            assert len(chapter_nodes) >= 1

        except ImportError:
            pytest.skip("python-docx not installed")


# =============================================================================
# Edge Cases
# =============================================================================


class TestEdgeCases:
    """Edge case tests."""

    def test_empty_content(self, parser: LlamaIndexLegalParser):
        chunks, tree = parser._create_hierarchical_chunks(
            content="",
            source_file="empty.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )
        assert len(chunks) == 0
        assert tree is not None
        assert tree.children == []

    def test_content_without_structure(self, parser: LlamaIndexLegalParser):
        """Content with no chapters/articles should produce no chunks."""
        chunks, tree = parser._create_hierarchical_chunks(
            content="This is just plain text without any legal structure.",
            source_file="plain.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )
        assert len(chunks) == 0
        assert tree.children == []

    def test_articles_without_chapters(self, parser: LlamaIndexLegalParser):
        content = """\
Điều 1. Phạm vi
Luật này quy định.
Điều 2. Đối tượng
Cá nhân, tổ chức.
"""
        chunks, tree = parser._create_hierarchical_chunks(
            content=content,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )
        assert len(chunks) == 2
        # Articles should be direct children of root (no chapters)
        assert len(tree.children) == 2
        assert tree.children[0].node_type == LegalNodeType.ARTICLE

    def test_single_article(self, parser: LlamaIndexLegalParser):
        content = "Điều 1. Điều duy nhất\nNội dung duy nhất."
        chunks, tree = parser._create_hierarchical_chunks(
            content=content,
            source_file="test.docx",
            law_id="X",
            law_name="Test",
            doc_kind="LAW",
        )
        assert len(chunks) == 1
        assert "DIEU=1" in chunks[0].chunk_id

    def test_normalize_text(self, parser: LlamaIndexLegalParser):
        text = "QUỐC HỘI\nCỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n\n\n\nNội dung"
        normalized = parser._normalize_text(text)
        assert "QUỐC HỘI" not in normalized
        assert "Nội dung" in normalized

    def test_extract_tables_markdown(self, parser: LlamaIndexLegalParser):
        content = """\
| STT | Nội dung |
| --- | --- |
| 1 | Điều kiện |
| 2 | Tiêu chuẩn |
"""
        tables = parser._extract_tables(content)
        assert len(tables) == 1
        assert tables[0]["num_rows"] == 2
        assert tables[0]["num_cols"] == 2
